from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt
from datetime import date
import pathlib
from docx.shared import Inches
## VARIÁVEIS 

equipamentos = {
    "EXT" : "Exaustor",
    "EXAUSTOR": "Exaustor",
    "BMB" : "Bomba",
    "BOMBA" : "Bomba",
    "FCL" : "Fan Coil",
    "VNT" : "Ventilador",
    "VENT" : "Ventilador",
    "VENTILADOR": "Ventilador",
    "CX" : "Ventilador",
    "CHILLER": "Compressor",
    "COMPRESSOR" : "Compressor",
    "FAN COIL": "Fan Coil",
    "BAGP" : "Bomba",
    "BAGS" : "Bomba"
}

defeitos = ["DESBALANCEAMENTO", "DESALINHAMENTO", "LUBRIFICAÇÃO", "ROLAMENTO", "BASE DANIFICADA", "FOLGA ", "FOLGAS"]

## FUNÇÕES

def WORD_arrumarAbreviacoes(tabela, index):
    for item in tabela.columns[index].cells:
        if item.text == "n" or item.text == "N":
            item.text = "Normal"
        if item.text == "A1" or item.text == "a1":
            item.text = "Aceitável"
        if item.text == "A2" or item.text == "a2":
            item.text = "Alerta"
        if item.text == "A3" or item.text == "a3":
            item.text = "Crítico"
        if item.text == "P" or item.text == "p":
            item.text = "Parado"
        WORD_formatarCelula(item)
                    

def WORD_arrumarOS(tabela, totLinhas):
    print(f"🛠 -> Adicionando contagem na coluna de OS...")
    countOS = 1
    for i in range(0, totLinhas):
        tabelaStatus = tabela.columns[4].cells[i]
        tabelaOS = tabela.columns[5].cells[i] 
        if tabelaStatus.text == "Aceitável" or tabelaStatus.text == "Alerta" or tabelaStatus.text == "Crítico":
            tabelaOS.text = WORD_arrumarCounts(countOS)
            WORD_formatarCelula(tabelaOS)
            countOS += 1
        else:
            tabelaOS.text = ""
    print(f"✔ -> Contagem adicionada na coluna de OS")

def WORD_arrumarCounts(count):
    if count < 10:
        return f"0{count}"
    else:
        return str(count)

def WORD_retornarData():
    data = str(date.today()).split("-")
    datacorreta = f"{data[2]}/{data[1]}/{data[0]}"
    return datacorreta
    
def WORD_arrumarTabelaOS_equipamento(documento, tabelasCount):
    print(f"\n🛠 -> Formatando tabelas de OS...")
    countOS = 1
    for i in range(2, tabelasCount):
        if i%2 == 0 :
            tabela = documento.tables[i]
            tabela.columns[3].cells[0].text = WORD_arrumarCounts(countOS)
            WORD_formatarCelula(tabela.columns[3].cells[0])
            countOS += 1
            WORD_arrumarAbreviacoes(tabela, 4)
    print(f"🛠 -> Tabelas de OS formatados!")

def WORD_deletarColuna(documento, table, columns):
    table = documento.tables[table]
    grid = table._tbl.find("w:tblGrid", table._tbl.nsmap)
    for cell in table.column_cells(columns):
        cell._tc.getparent().remove(cell._tc)
    col_elem = grid[columns]
    grid.remove(col_elem)

def WORD_formatarCelula(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Arial"

def WORD_formatarCabecalho(celula):
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Arial"
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#A6A6A6")
    table_cell_properties.append(shade_obj)    

def WORD_formatarData(celula):
    celula.text = WORD_retornarData()
    for paragraph in celula.paragraphs:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            run.font.bold = True
            run.font.name = "Aharoni"
            run.font.size = Pt(36)
    cell_xml_element = celula._tc
    table_cell_properties = cell_xml_element.get_or_add_tcPr()
    shade_obj = OxmlElement('w:shd')
    shade_obj.set(qn('w:fill'), "#FFF2CC")
    table_cell_properties.append(shade_obj)    

def WORD_addCabecalhoVertical(tabela, totLinhas):
    for i in range(0, totLinhas):
        WORD_formatarCabecalho(tabela.columns[0].cells[i])

def WORD_colunaValores(tabela, index):
    valores = []
    for item in tabela.columns[index].cells:
        valores.append(item.text)
    return valores

def WORD_arrumarEquipamentoTabela(tabela, totLinhas):
    print(f"🛠 -> Arrumando conjuntos na tabela de listagem...")
    for i in range(0,totLinhas): 
        texto = tabela.columns[1].cells[i].text
        for chave, valor in equipamentos.items():
            if chave in texto.upper():
                tabela.columns[2].cells[i].text = valor
            else: 
                texto.split(" ")[0].capitalize()
        WORD_formatarCelula(tabela.columns[2].cells[i])
    print(f"✔ -> Conjuntos na tabela de listagem arrumados!")


def WORD_indentificarDefeito(documento, tabelasCount):
    print(f"🛠 -> Identificando defeitos das OS...")
    statusArray = []
    defeitoArray = []

    for i in range(2, tabelasCount):
        if i%2 == 0:
            tabelaStatus = documento.tables[i]
            statusArray.append(tabelaStatus.columns[4].cells[2].text)
    
    for i in range(2, tabelasCount):
        if i%2 != 0:
            tabelaDefeito = documento.tables[i]
            texto = tabelaDefeito.columns[0].cells[1].text.upper()
            arraytemp = []
            for i in defeitos:
                if i in texto:
                    arraytemp.append("FOLGAS" if i == "FOLGA " else i)
            if len(arraytemp) == 0:
                arraytemp.append("OUTROS")
            defeitoArray.append(arraytemp)
            arraytemp = []
    
    arrayStatusDefeitos = []
    for i in range(len(statusArray)):
        arrayStatusDefeitos.append([statusArray[i], defeitoArray[i]])

    print(f"✔ -> Defeitos identificados! (Enviando Excel)")
    return arrayStatusDefeitos

def WORD_addGraficos(paragrafo, nm):
    paragrafo.text = ''
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img = paragrafo.add_run()
    img.add_picture(rf"{str(pathlib.Path().resolve())}\chart{nm}.png", width=Inches(4))

