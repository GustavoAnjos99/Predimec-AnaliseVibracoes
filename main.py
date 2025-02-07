from docx import Document
from openpyxl import load_workbook
from functions_WORD import *
from functions_EXCEL import *
import os
import pathlib
from win32com.client import Dispatch
import sys
import time
import pathlib

## Definição das Funções ===========
def excluirImagensPATH(imagem):
    if os.path.exists(imagem):
        os.remove(imagem)

def pegarGraficosExcel(app, workbook_file_name, workbook):
    app.DisplayAlerts = False

    for i, sheet in enumerate(workbook.Worksheets):
        for chartObject in sheet.ChartObjects():
            chartObject.Chart.Export(rf"{str(pathlib.Path().resolve())}\chart{str(i+1)}.png")
            i +=1
    workbook.Close(SaveChanges=False, Filename=workbook_file_name)

## Inicialização do app ==========
print(r"""
________            _____________                           __________  
___  __ \_________________  /__(_)______ _______________    ___  ___  \ 
__  /_/ /_  ___/  _ \  __  /__  /__  __ `__ \  _ \  ___/    __  / _ \  |
_  ____/_  /   /  __/ /_/ / _  / _  / / / / /  __/ /__      _  / , _/ / 
/_/     /_/    \___/\__,_/  /_/  /_/ /_/ /_/\___/\___/      | /_/|_| /  
                                                             \______/   
Iniciando processo de formatação...
      """)

ARQUIVO_WORD = ''
ARQUIVO_EXCEL = ''

try: 
    arquivos = os.listdir('./')
    for arquivo in arquivos:
        if arquivo.endswith(".docx"):
            ARQUIVO_WORD = arquivo
        if arquivo.endswith(".xlsm") or arquivo.endswith(".xlsx"):
            ARQUIVO_EXCEL = arquivo
    f = open(ARQUIVO_WORD, 'rb')
    g = open(ARQUIVO_EXCEL, 'rb')
    documentoWord = Document(f)
    documentoExcel = load_workbook(g)
except:
    print("ERRO: Erro ao identificar arquivos para formatação.")
    time.sleep(10)
    sys.exit(1)
    
arquivocorreto = False
while arquivocorreto == False:
    padraoArquivo = "<>:\"/|?*"
    nomeArquivoUser = str(input("\n💾 -> Digite o nome dos arquivos a serem salvos [EXCEL e WORD]: "))
    for i in padraoArquivo:
        if i in nomeArquivoUser:
            print("\n⛔ -> O nome do arquivo NÃO pode ter os seguintes caracteres: '/' , '<', '>', '\', '|', '?', '*'.")
            arquivocorreto = False
            break
        else: 
            arquivocorreto = True

## WORD ==========
tabela = documentoWord.tables[1]
totLinhas = len(tabela.rows)
totColunas = int(len(tabela.columns) / 2)
tabelasCount = len(documentoWord.tables)


WORD_formatarData(documentoWord.tables[0].columns[0].cells[0])
WORD_deletarColuna(documentoWord,1,0)
WORD_deletarColuna(documentoWord,1,2)
WORD_deletarColuna(documentoWord,1,6)
WORD_deletarColuna(documentoWord,1,6)
WORD_arrumarTabelaOS_equipamento(documentoWord, tabelasCount)
WORD_arrumarEquipamentoTabela(tabela, totLinhas)
WORD_addCabecalhoVertical(tabela, totLinhas)

print(f"🛠 -> Formatando status na tabela de listagem...")
for i in range(0, totColunas):
    WORD_arrumarAbreviacoes(tabela, i)
print(f"✔ -> Status na tabela de listagem formatados!")

WORD_arrumarOS(tabela, totLinhas)

## EXCEL ==========
documentoExcel.active = documentoExcel['Listagem']
planilhaListagem = documentoExcel.active
planilhaGraficos = documentoExcel['Gráficos']

addColunaListagem(WORD_colunaValores(tabela, 4), planilhaListagem, planilhaGraficos)
EXCEL_arrumarTabela_2(planilhaListagem, WORD_indentificarDefeito(documentoWord, tabelasCount))
EXCEL_arrumarTabela_3(planilhaGraficos)
EXCEL_corrigirFormulas(planilhaGraficos)

## SALVAMENTO DO ARQUIVO ========
PASTA_RESULTADOS = "RELATÓRIOS FORMATADOS"
os.makedirs(PASTA_RESULTADOS, exist_ok=True)

caminhoWord = os.path.join(PASTA_RESULTADOS, f"{nomeArquivoUser}.docx")
caminhoExcel = os.path.join(PASTA_RESULTADOS, f"{nomeArquivoUser}.xlsx")
documentoExcel.save(caminhoExcel)
try:
    app = Dispatch("Excel.Application")
    workbook_file_name = rf"{str(pathlib.Path().resolve())}\RELATÓRIOS FORMATADOS\{nomeArquivoUser}"
    workbook = app.Workbooks.Open(Filename=workbook_file_name)
    pegarGraficosExcel(app, workbook_file_name, workbook)

    for i in documentoWord.paragraphs:
        if i.text == "[gráfico1]":
            WORD_addGraficos(i, 2)
        if i.text == "[gráfico2]":
            WORD_addGraficos(i, 3)
        if i.text == "[gráfico3]":
            WORD_addGraficos(i, 4)
    print("✔ -> Gráficos adicionados no arquivo WORD!")
except:
    print("\nERRO: Erro ao inserir imagens dos gráficos do arquivo WORD.\n💡 -> Certifique-se de que quando salvou o arquivo excel as imagens estavam dentro dos limites da tela.")
    excluirImagensPATH("chart1.png")
    excluirImagensPATH("chart2.png")
    excluirImagensPATH("chart3.png")
    excluirImagensPATH("chart4.png")
    time.sleep(10)
    sys.exit(1)

documentoWord.save(caminhoWord)

excluirImagensPATH("chart1.png")
excluirImagensPATH("chart2.png")
excluirImagensPATH("chart3.png")
excluirImagensPATH("chart4.png")

print("\nArquivos formatados com sucesso!\n")
time.sleep(10)
